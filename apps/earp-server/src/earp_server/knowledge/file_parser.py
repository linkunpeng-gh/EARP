"""Document text extraction for uploads — .docx / .pdf / plain text.

Word (.docx) is a zip container; python-docx extracts paragraph runs.
PDF uses pypdf page extraction (digital text only — scanned/OCR pages yield
empty text and are rejected with a clear error).

Plain-text extensions (txt/md/csv/json/html) are passed through as-is.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".html", ".htm", ".log"}


class FileParseError(Exception):
    """Raised when a file cannot be parsed to text."""


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded file based on its extension."""
    ext = _ext(filename)
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext in TEXT_EXTENSIONS:
        for enc in ("utf-8", "gb18030", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        raise FileParseError(f"无法解码文本文件 {filename}（尝试 utf-8/gb18030 均失败）")
    raise FileParseError(f"不支持的文件类型 {ext or '(无扩展名)'}（支持: docx/pdf/txt/md/csv/json/html）")


def _ext(filename: str) -> str:
    import os

    return os.path.splitext(filename or "")[1].lower()


def _parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("python-docx 未安装，无法解析 Word 文档") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise FileParseError(f"Word 文档解析失败（文件可能损坏或不是 .docx 格式）: {exc}") from exc

    parts: list[str] = []
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # tables: each row joined by tab, rows separated by newline
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    text = "\n\n".join(parts)
    if not text.strip():
        raise FileParseError("Word 文档未提取到任何文本（可能是纯图片文档）")
    return text


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("pypdf 未安装，无法解析 PDF") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for i, page in enumerate(reader.pages, 1):
            t = (page.extract_text() or "").strip()
            if t:
                pages.append(f"--- 第 {i} 页 ---\n{t}")
    except Exception as exc:
        raise FileParseError(f"PDF 解析失败（文件可能损坏或加密）: {exc}") from exc

    text = "\n\n".join(pages)
    if not text.strip():
        raise FileParseError("PDF 未提取到任何文本（可能是扫描件，需要 OCR，暂不支持）")
    return text
